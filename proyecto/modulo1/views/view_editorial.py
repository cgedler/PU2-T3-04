import os
from pathlib import Path
from datetime import datetime
from django import forms
from django.shortcuts import get_object_or_404, render
from django.http import HttpResponse
from datetime import date
from django.utils import timezone
from django.views.generic import TemplateView, View
from django.views.generic import ListView, DetailView
from django.views.generic.edit import CreateView, UpdateView, DeleteView
from django.utils.decorators import method_decorator
from django.contrib.auth.decorators import login_required
from django.urls import reverse
from django.contrib import messages
from django.contrib.messages.views import SuccessMessageMixin
from django.db.models import Q
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from openpyxl import Workbook

from modulo1.models.model_base import Editorial
from modulo1.forms.form_editorial import EditorialForm
from modulo1.reportes.generar_pdf import *


BASE_DIR = Path(__file__).resolve().parent.parent

class CrearEditorial(SuccessMessageMixin, CreateView):
    model = Editorial
    form_class = EditorialForm
    template_name = "editoriales/create.html"
    success_message = "Creado Correctamente!"

    def get_success_url(self):
        return reverse("modulo1:EditorialListado")

class ListadoEditorial(ListView):
    model = Editorial
    template_name = "editoriales/list.html"
    queryset = Editorial.objects.all().order_by('-created')
    context_object_name = 'Editorial'
    paginate_by = 1000


class DetalleEditorial(DetailView):
    model = Editorial
    form_class = EditorialForm
    template_name = "editoriales/detail.html"


class ActualizarEditorial(SuccessMessageMixin, UpdateView):
    model = Editorial
    form_class = EditorialForm
    template_name = "editoriales/update.html"
    success_message = "Actualizado Correctamente!"

    def get_success_url(self):
        return reverse("modulo1:EditorialListado")


class EliminarEditorial(SuccessMessageMixin, DeleteView):
    model = Editorial

    def get_success_url(self):
        success_message = "Eliminado Correctamente!"
        messages.success(self.request, (success_message))
        return reverse("modulo1:EditorialListado")


class GeneratePdfListEditorial(View):
    def get(self, request, *args, **kwargs):
        now = timezone.now()
        data = {
            "section_title": "Listado",
            "module_title": "Editorial",
            "today": now,
            "empresa_nombre": "Nombre Empresa",
            "empresa_rif": "RIF",
            "empresa_direstado": "Estado",
            "empresa_dirciudad": "Ciudad",
            "empresa_direccion": "Direccion...",
            "empresa_postal": "Codigo Postal",
            "context_pdf": Editorial.objects.all()
        }
        pdf = render_to_pdf("editoriales/listadopdf.html", data)
        return HttpResponse(pdf, content_type="application/pdf")


class GeneratePdfEditorial(View):
    def get(self, request, *args, **kwargs):
        id = kwargs["pk"]
        object_list = Editorial.objects.get(id=id)
        now = timezone.now()
        data = {
            "section_title": "Listado",
            "module_title": "Editorial",
            "today": now,
            "empresa_nombre": "Nombre Empresa",
            "empresa_rif": "RIF",
            "empresa_direstado": "Estado",
            "empresa_dirciudad": "Ciudad",
            "empresa_direccion": "Direccion...",
            "empresa_postal": "Codigo Postal",
            "context_pdf": object_list
        }
        pdf = render_to_pdf("editoriales/Editorialpdf.html", data)
        return HttpResponse(pdf, content_type="application/pdf")
    

class GenerateWordEditorial(SuccessMessageMixin, View):
    def get(self, request, *args, **kwargs):
        id = str
        id = kwargs["pk"]
        object_list = Editorial.objects.get(id=id)
        document = Document(os.path.join(BASE_DIR, 'reportes/base.docx')) # Utilizar una plantilla
        #document = Document() # Crear nuevo documento de Word
        style = document.styles['Normal']
        sections = document.sections
        for section in sections:
            section.top_margin = Cm(1.0)
            section.bottom_margin = Cm(1.0)
            section.left_margin = Cm(1.0)
            section.right_margin = Cm(1.0)
        font = style.font
        font.name = 'Arial' 
        font.size = Pt(9)
        
        document.add_paragraph("ID: " + str(object_list.id))
        document.add_paragraph("Nombre: " + object_list.nombre)
        document.add_paragraph("Direccion: " + object_list.direccion)
        document.add_paragraph("Telefono: " + object_list.telefono)
        document.add_paragraph("Contacto: " + object_list.contacto)
        document.add_paragraph("Web : " + object_list.web)
  
        document.add_page_break()
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename={date}-Editorial.docx'.format(
            date=datetime.now().strftime('%d-%m-%Y'),
        )
        document.save(response)
        return response
    
class GenerateExcelEditorial(SuccessMessageMixin, View):
    def get(self, request, *args, **kwargs):
        id = str
        id = kwargs["pk"]
        object_list = Editorial.objects.get(id=id)
        document = Workbook() # Crear nuevo documento de Excel

        hoja = document.active
        
        hoja["A1"] = "ID: " 
        hoja["B1"] = str(object_list.id)
        hoja["A2"] = "Nombre: "
        hoja["B2"] = object_list.nombre
        hoja["A3"] = "Direccion: "
        hoja["B3"] = object_list.direccion
        hoja["A4"] = "Telefono: "
        hoja["B4"] = object_list.telefono
        hoja["A5"] = "Contacto: "
        hoja["B5"] = object_list.contacto
        hoja["A6"] = "Web "
        hoja["B6"] = object_list.web
  
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        )
        response['Content-Disposition'] = 'attachment; filename={date}-Editorial.xlsx'.format(
            date=datetime.now().strftime('%d-%m-%Y'),
        )
        document.save(response)
        return response