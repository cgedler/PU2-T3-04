import os
from pathlib import Path
from datetime import datetime
from django import forms
from django.shortcuts import get_object_or_404, render
from django.http import HttpResponse
from datetime import date
from django.utils import timezone
from django.views.generic import TemplateView, View
from django.views.generic import ListView, DetailView
from django.views.generic.edit import CreateView, UpdateView, DeleteView
from django.utils.decorators import method_decorator
from django.contrib.auth.decorators import login_required
from django.urls import reverse
from django.contrib import messages
from django.contrib.messages.views import SuccessMessageMixin
from django.db.models import Q
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from openpyxl import Workbook

from modulo1.models.model_base import Libro
from modulo1.forms.form_libro import LibroForm
from modulo1.reportes.generar_pdf import *


BASE_DIR = Path(__file__).resolve().parent.parent

class CrearLibro(SuccessMessageMixin, CreateView):
    model = Libro
    form_class = LibroForm
    template_name = 'libros/create.html'
    success_message = 'Creado Correctamente!'

    def get_success_url(self):
        return reverse('modulo1:LibroListado')

class ListadoLibro(ListView):
    model = Libro
    template_name = 'libros/list.html'
    queryset = Libro.objects.all().order_by('-created')
    context_object_name = 'Libro'
    paginate_by = 1000


class DetalleLibro(DetailView):
    model = Libro
    form_class = LibroForm
    template_name = 'libros/detail.html'


class ActualizarLibro(SuccessMessageMixin, UpdateView):
    model = Libro
    form_class = LibroForm
    template_name = 'libros/update.html'
    success_message = 'Actualizado Correctamente!'

    def get_success_url(self):
        return reverse('modulo1:LibroListado')


class EliminarLibro(SuccessMessageMixin, DeleteView):
    model = Libro

    def get_success_url(self):
        success_message = 'Eliminado Correctamente!'
        messages.success(self.request, (success_message))
        return reverse('modulo1:LibroListado')


class SearchResultsLibros(ListView):
    model = Libro
    template_name = 'libros/filtro.html'

    def get_queryset(self):
        query = self.request.GET.get('q')
        object_list = Libro.objects.filter(
            Q(isbn__icontains=query) | Q(titulo__icontains=query)
        )
        return object_list


class GeneratePdfListLibros(View):
    def get(self, request, *args, **kwargs):
        now = timezone.now()
        data = {
            'section_title': 'Listado',
            'module_title': 'Libro',
            'today': now,
            'empresa_nombre': 'Nombre Empresa',
            'empresa_rif': 'RIF',
            'empresa_direstado': 'Estado',
            'empresa_dirciudad': 'Ciudad',
            'empresa_direccion': 'Direccion...',
            'empresa_postal': 'Codigo Postal',
            'context_pdf': Libro.objects.all()
        }
        pdf = render_to_pdf('libros/listadopdf.html', data)
        return HttpResponse(pdf, content_type='application/pdf')


class GeneratePdfLibro(View):
    def get(self, request, *args, **kwargs):
        id = kwargs['pk']
        object_list = Libro.objects.get(id=id)
        now = timezone.now()
        data = {
            'section_title': 'Listado',
            'module_title': 'Libro',
            'today': now,
            'empresa_nombre': 'Nombre Empresa',
            'empresa_rif': 'RIF',
            'empresa_direstado': 'Estado',
            'empresa_dirciudad': 'Ciudad',
            'empresa_direccion': 'Direccion...',
            'empresa_postal': 'Codigo Postal',
            'context_pdf': object_list
        }
        pdf = render_to_pdf('libros/libropdf.html', data)
        return HttpResponse(pdf, content_type='application/pdf')
    

class GenerateWordLibro(SuccessMessageMixin, View):
    def get(self, request, *args, **kwargs):
        id = str
        id = kwargs['pk']
        object_list = Libro.objects.get(id=id)
        document = Document(os.path.join(BASE_DIR, 'reportes/base.docx')) # Utilizar una plantilla
        #document = Document() # Crear nuevo documento de Word
        style = document.styles['Normal']
        sections = document.sections
        for section in sections:
            section.top_margin = Cm(1.0)
            section.bottom_margin = Cm(1.0)
            section.left_margin = Cm(1.0)
            section.right_margin = Cm(1.0)
        font = style.font
        font.name = 'Arial' 
        font.size = Pt(9)
        
        document.add_paragraph('ID: ' + str(object_list.id))
        document.add_paragraph('ISBN: ' + object_list.isbn)
        document.add_paragraph('Titulo: ' + object_list.titulo)
        document.add_paragraph('Autor: ' + object_list.autor)
        document.add_paragraph('Fecha de Publicacion: ' + str(object_list.publicacion))
        document.add_paragraph('Paginas del Libro: ' + str(object_list.paginas))
        document.add_paragraph('Precio del Libro: ' + str(object_list.precio))
        document.add_paragraph('Distribuidor : ' + str(object_list.id_distribuidora))
        document.add_paragraph('Editorial: ' + str(object_list.id_editorial))
  
        document.add_page_break()
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename={date}-Libro.docx'.format(
            date=datetime.now().strftime('%d-%m-%Y'),
        )
        document.save(response)
        return response
    
class GenerateExcelLibro(SuccessMessageMixin, View):
    def get(self, request, *args, **kwargs):
        id = str
        id = kwargs['pk']
        object_list = Libro.objects.get(id=id)
        document = Workbook() # Crear nuevo documento de Excel

        hoja = document.active
        
        hoja['A1'] = 'ID: ' 
        hoja['B1'] = str(object_list.id)
        hoja['A2'] = 'ISBN: '
        hoja['B2'] = object_list.isbn
        hoja['A3'] = 'Titulo: '
        hoja['B3'] = object_list.titulo
        hoja['A4'] = 'Autor: '
        hoja['B4'] = object_list.autor
        hoja['A5'] = 'Fecha de Publicacion: '
        hoja['B5'] = str(object_list.publicacion)
        hoja['A6'] = 'Paginas del Libro: '
        hoja['B6'] = str(object_list.paginas)
        hoja['A7'] = 'Precio del Libro: '
        hoja['B7'] = str(object_list.precio)
        hoja['A8'] = 'Distribuidor : '
        hoja['B8'] = str(object_list.id_distribuidora)
        hoja['A9'] = 'Editorial: '
        hoja['B9'] = str(object_list.id_editorial)
  
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        )
        response['Content-Disposition'] = 'attachment; filename={date}-Libro.xlsx'.format(
            date=datetime.now().strftime('%d-%m-%Y'),
        )
        document.save(response)
        return response


def CountLibros(valor):
    query = Libro.objects.filter(Q(precio__gte=valor))
    total = query.count()
    return total


def GraficoLibro(request):    
    labels = ['5$', '10$', '15$', '20$', '30$', '50', '100$', '150$', '200$']
    L_5 = CountLibros(5)
    L_10 = CountLibros(10)
    L_15 = CountLibros(15)
    L_20 = CountLibros(20)
    L_30 = CountLibros(30)
    L_50 = CountLibros(50)
    L_100 = CountLibros(100)
    L_150 = CountLibros(150)
    L_200 = CountLibros(200)

    data = [L_5, L_10, L_15, L_20, L_30, L_50, L_100, L_150, L_200]
    
    return render(request, 'libros/grafico.html',  {
        'labels': labels,
        'data': data,
        })